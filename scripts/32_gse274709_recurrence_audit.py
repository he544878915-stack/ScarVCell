from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / ".stage" / "public_expansion" / "GSE274709_PMC12337814_supplementary" / "PATH-267-105-s001.docx"
TABLES = ROOT / "results" / "tables"
TABLES.mkdir(parents=True, exist_ok=True)


def main() -> None:
    document = Document(SOURCE)
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    table_records = []
    for table_index, table in enumerate(document.tables, 1):
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        table_records.append({"table_index": table_index, "rows": rows})
    relevant_paragraphs = [
        text for text in paragraphs if re.search(r"recurr|KL\s*\d|keloid\s*\d|Table\s*S1", text, flags=re.IGNORECASE)
    ]
    relevant_tables = []
    for table in table_records:
        flattened = " ".join(cell for row in table["rows"] for cell in row)
        if re.search(r"recurr|KL\s*\d|keloid\s*\d|PIEZO2", flattened, flags=re.IGNORECASE):
            relevant_tables.append(table)
    payload = {
        "source": str(SOURCE.relative_to(ROOT)),
        "paragraphs": len(paragraphs),
        "tables": len(table_records),
        "relevant_paragraphs": relevant_paragraphs,
        "relevant_tables": relevant_tables,
    }
    output = TABLES / "GSE274709_recurrence_mapping_audit.json"
    output.write_text(json.dumps(payload, indent=2), encoding="utf-8")
    print(json.dumps(payload, indent=2))


if __name__ == "__main__":
    main()
